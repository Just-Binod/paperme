from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_question_paper_doc(data, questions_text):
    """
    Create a formatted Word document for the question paper
    following Pokhara University format.
    """
    try:
        doc = Document()
        
        # Set default font for entire document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # University Name
        uni = doc.add_paragraph("POKHARA UNIVERSITY")
        uni.runs[0].bold = True
        uni.runs[0].font.size = Pt(14)
        uni.runs[0].font.name = 'Times New Roman'
        uni.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer

        # College Name
        college = doc.add_paragraph(data.get('college_name', ''))
        college.runs[0].bold = True
        college.runs[0].font.size = Pt(16)
        college.runs[0].font.name = 'Times New Roman'
        college.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Address
        if data.get('college_address'):
            addr = doc.add_paragraph(data['college_address'])
            addr.runs[0].font.name = 'Times New Roman'
            addr.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer

        # Details
        details = doc.add_paragraph()
        details.add_run(f"Level: Bachelor    Semester: Spring    Year: {data.get('year', '')}\n")
        details.add_run(f"Programme: {data.get('program', '')}    Full Marks: {data.get('full_marks', '')}\n")
        details.add_run(f"Course: {data.get('course_name', '')}    Pass Marks: {data.get('pass_marks', '')}\n")
        details.add_run(f"Time: {data.get('time_hours', '')} hrs.")
        
        for run in details.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        details.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer

        # Instructions
        inst = doc.add_paragraph()
        inst.add_run("Candidates are required to give their answers in their own words as far as practicable.\n")
        inst.add_run("The figures in the margin indicate full marks.\n")
        inst.add_run("Attempt all the questions.")
        
        for run in inst.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        doc.add_paragraph()  # spacer before questions

        # Add questions
        for line in questions_text.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            p = doc.add_paragraph(line)
            
            # Format question numbers
            if line and line[0].isdigit() and '.' in line[:3]:
                p.runs[0].bold = True
            
            # Indent sub-questions
            elif line.startswith(('a)', 'b)', 'OR', 'i)', 'ii)', 'iii)')):
                p.paragraph_format.left_indent = Inches(0.5)
            
            # Set font for all runs
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        return doc
        
    except Exception as e:
        raise Exception(f"Document creation failed: {str(e)}")





#############

##########

##########



# from docx import Document
# from docx.shared import Pt, Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# def create_question_paper_doc(data, questions_text):
#     doc = Document()





#     # University Name
#     uni = doc.add_paragraph("POKHARA UNIVERSITY")
#     uni.runs[0].bold = True
#     uni.runs[0].font.size = Pt(12)
#     uni.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     doc.add_paragraph()  # spacer

#     # Details table (single row or plain text)
#     details = doc.add_paragraph()
#     details.add_run(f"Level: Bachelor    Semester: Spring    Year: {data['year']}\n")
#     details.add_run(f"Programme: {data['program']}    Full Marks: {data['full_marks']}\n")
#     details.add_run(f"Course: {data['course_name']}    Pass Marks: {data['pass_marks']}\n")
#     details.add_run(f"Time: {data['time_hours']} hrs.")
#     details.alignment = WD_ALIGN_PARAGRAPH.CENTER   

#     # College Name
#     college = doc.add_paragraph(data['college_name'])
#     college.runs[0].bold = True
#     college.runs[0].font.size = Pt(16)
#     college.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     # Address
#     addr = doc.add_paragraph(data['college_address'])
#     addr.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     # # Exam Title
#     # title = doc.add_paragraph(data['exam_type'])
#     # title.runs[0].bold = True
#     # title.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     # # Table for details
#     # table = doc.add_table(rows=1, cols=2)
#     # table.style = 'Table Grid'
#     # row = table.rows[0].cells
#     # row[0].text = f"Level: Bachelor\nProgram: {data['program']}\nCourse: {data['course_name']}\nSemester: {data['semester']}"
#     # row[1].text = f"Year: {data['year']}\nFull Marks: {data['full_marks']}\nPass Marks: {data['pass_marks']}\nTime: {data['time_hours']} hrs."

#     doc.add_paragraph()

#     # Instructions
#     inst = doc.add_paragraph("Candidates are required to give their answers in their own words as far as practicable.\n")
#     inst.add_run("The figures in the margin indicate full marks.\n")
#     inst.add_run("Attempt all the questions.")

#     # doc.add_page_break()

#     # Add questions
#     for line in questions_text.split('\n'):
#         line = line.strip()
#         if not line:
#             continue
#         p = doc.add_paragraph(line)
#         # if line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.')):
#             # p.runs[0].bold = True
#             # p.runs[0].font.size = Pt(12)
#             # p.runs[1].font.size = Pt(12)
#         # elif line.startswith(('a)', 'b)', 'OR')):
#         #     p.paragraph_format.left_indent = Inches(0.5)

#     return doc